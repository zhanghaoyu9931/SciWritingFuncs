
import os
import sys
import pandas as pd
# for docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from typing import Optional, Union


def df_to_three_line_table(
    df,
    output_path: str = "three_line_table.docx",
    table_title: str = "Table 1. example",
    header_bold: bool = True,
    top_border_width: int = 12,      # width of top border in 1/8 pt units
    middle_border_width: int = 4,   # width of middle border in 1/8 pt units
    bottom_border_width: int = 12,   # width of bottom border in 1/8 pt units
    font_name: Optional[str] = None,
    font_size: Optional[Union[int, float]] = 8,
    align_header: str = "center",   # header alignment option
    align_data: str = "center",     # data alignment option
    num_precision = 4,
    cols_anno_best = {'F1_score': 'max'} # annotate the best values in the table
) -> None:
    """
    convert a dataframe to three-line table in word format
    
    args:
        df: the input dataframe
        output_path: output word document path
        header_bold: enable bold header text
        top_border_width: top line thickness
        middle_border_width: middle line thickness
        bottom_border_width: bottom line thickness
        font_name: specify font family
        font_size: specify font size in points
        align_header: header text alignment
        align_data: data cell alignment
    """
    try:
        # === title settings ===
        doc = Document()
        title_para = doc.add_heading(level=1)
        title_run = title_para.add_run(table_title)
        title_run.bold = True
        title_run.font.size = Pt(10)
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        title_run.font.name = 'Arial'
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        
        # === table settings ===
        table = doc.add_table(rows=len(df)+1, cols=len(df.columns),)# style='')
        # === Core: Three-line border settings ===
        def set_cell_border(cell, border_settings):
            """Set cell borders using OOXML specifications"""
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.first_child_found_in("w:tcBorders")
            if tcBorders is None:
                tcBorders = OxmlElement("w:tcBorders")
                tcPr.append(tcBorders)
            
            # clear existing borders
            for edge in ['top', 'start', 'left', 'bottom', 'end', 'right', 'insideH', 'insideV']:
                edge_tag = OxmlElement(f"w:{edge}")
                edge_tag.set(qn('w:val'), 'nil')
                edge_tag.set(qn('w:sz'), '0')
                edge_tag.set(qn('w:color'), "auto")
                tcBorders.append(edge_tag)
            
            # set new borders
            for edge, settings in border_settings.items():
                edge_tag = OxmlElement(f"w:{edge}")
                for key, value in settings.items():
                    edge_tag.set(qn(f'w:{key}'), str(value))
                tcBorders.append(edge_tag)

        # remove all borders
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell, {})

        # 1. set top line and set middle line
        for cell in table.rows[0].cells:
            set_cell_border(cell, {
                'top': {'val': 'single', 'sz': top_border_width, 'color': '000000'},
                'bottom': {'val': 'single', 'sz': middle_border_width, 'color': '000000'}
            })
        # 2. set bottom line
        for cell in table.rows[-1].cells:
            set_cell_border(cell, {
                'bottom': {'val': 'single', 'sz': bottom_border_width, 'color': '000000'}
            })

        # === fill values ===
        cells2annotate = []
        # header
        for col_idx, header in enumerate(df.columns):
            cell = table.cell(0, col_idx)
            cell.text = str(header)
            if header_bold:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            cell.paragraphs[0].alignment = getattr(WD_ALIGN_PARAGRAPH, align_header.upper())
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # set font size
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(font_size)
                if font_name:
                    run.font.name = font_name
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            # get the best cell to annotate
            if header in cols_anno_best:
                best_met = cols_anno_best[header]
                try:
                    df[header] = df[header].astype(float)
                    best_row_idx = eval(f"df[header].idx{best_met}()")
                    cells2annotate.append((best_row_idx+1, col_idx))
                except:
                    pass

        # body
        for row_idx, row_data in enumerate(df.values, start=1):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                try:
                    # set number format
                    cell_data = float(cell_data)
                    if isinstance(cell_data, float):
                        cell_data = f"{cell_data:.{num_precision}f}"
                except:
                    pass
                
                cell.text = str(cell_data)
                cell.paragraphs[0].alignment = getattr(WD_ALIGN_PARAGRAPH, align_data.upper())
                # set font size
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(font_size)
                    if font_name:
                        run.font.name = font_name
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    if (row_idx, col_idx) in cells2annotate:
                        run.bold = True

        doc.save(output_path)
        print(f"Save the table in: {output_path}")

    except Exception as e:
        print(f"Error: {str(e)}", file=sys.stderr)
        raise

import cv2
import pytesseract
import pandas as pd
from PIL import Image
def image_table_to_csv(image_path, output_csv):
    img = cv2.imread(image_path)
    # to gray figure
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
    
    # OCR processing
    text = pytesseract.image_to_string(thresh, config='--psm 6')
    lines = [line.split() for line in text.split('\n') if line.strip()]
    
    # clean data and save csv
    header = lines[0]
    col_n = len(header)
    data = lines[1:]
    data = [x if len(x) == col_n else (['_'.join(x[:len(x)-col_n+1])] + x[len(x)-col_n+1:]) for x in data]
    
    df = pd.DataFrame(data, columns=header)
    df.to_csv(output_csv, index=False)
    return df
